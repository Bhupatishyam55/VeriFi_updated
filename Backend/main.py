# main.py
import io
import uuid
import re
import docx
import os
import json
import faiss
from datetime import datetime
from typing import Dict, List, Optional, Any
from contextlib import asynccontextmanager # Added for lifespan management

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import pdfplumber
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
from pypdf import PdfReader

# Importing your custom logic
from vector_store import search_duplicate, add_to_index
from fraud_detection import detect_pii, analyze_metadata, extract_advanced_entities 
from image_forensics import detect_tampering, get_image_phash
from pydantic import BaseModel, Field

# ----- SYSTEM RESET LOGIC (New) ----- #

def reset_system_data():
    """Wipes the FAISS index, the hash JSON, and the in-memory database."""
    
    # 1. Clear the in-memory DB
    global db
    db.clear() 
    
    # 2. Reset hash.json
    try:
        with open("hash.json", "w") as f:
            json.dump({}, f)
    except Exception as e:
        print(f"Error resetting hash.json: {e}")
        
    # 3. Reset docs.index (FAISS)
    try:
        dimension = 384 # Standard for all-MiniLM-L6-v2
        index = faiss.IndexFlatIP(dimension)
        faiss.write_index(index, "docs.index")
    except Exception as e:
        print(f"Error resetting FAISS index: {e}")
        
    return True

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Modern replacement for startup events to manage system resets on boot."""
    # This runs every time the server starts (Render boot/restart)
    reset_system_data()
    yield

# ----- Application Initialization ----- #

app = FastAPI(
    title="AP FraudShield Final", 
    version="1.0.0",
    lifespan=lifespan # Added lifespan to the app
)

# Get allowed origins from environment or use defaults
ALLOWED_ORIGINS = os.getenv(
    "CORS_ORIGINS",
    "http://localhost:3000,https://localhost:3000"
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

db = {}

# ----- Schemas (Keep existing) ----- #

class AnomalyItem(BaseModel):
    type: str
    description: str
    confidence: float = Field(..., ge=0, le=1)

class ScanResult(BaseModel):
    file_id: str
    filename: str
    file_url: str 
    fraud_score: int
    severity: str
    anomalies: List[AnomalyItem]
    text_content: str
    extracted_tables: List
    processing_time: int
    confidence: float = Field(..., ge=0, le=1)

class UploadResponse(BaseModel):
    task_id: str
    message: str

class AlertRequest(BaseModel):
    message: str

class AlertResponse(BaseModel):
    status: str

# ----- Helpers (Keep existing) ----- #

def clean_text(text: str) -> str:
    cleaned = re.sub(r"[^\w\s.,:/-]", " ", text)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()

def extract_text_from_file(content: bytes, filename: str) -> str:
    name = filename.lower()
    # Logic remains identical to your original provided file
    if name.endswith((".jpg", ".jpeg", ".png")):
        try:
            image = Image.open(io.BytesIO(content)).convert("L")
            return clean_text(pytesseract.image_to_string(image))
        except Exception: return ""

    if name.endswith(".docx") or name.endswith(".doc"):
        try:
            doc = docx.Document(io.BytesIO(content))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip(): full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip(): full_text.append(cell.text)
            return clean_text(" ".join(full_text))
        except Exception: return ""

    if name.endswith(".pdf"):
        fast_text = ""
        try:
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                if i >= 5: break
                pages_text.append(page.extract_text() or "")
            fast_text = clean_text(" ".join(pages_text))
            if len(fast_text) > 50: return fast_text
        except Exception: pass

        try:
            images = convert_from_bytes(content, dpi=150, first_page=1, last_page=3)
            ocr_text = " ".join(pytesseract.image_to_string(img.convert("L")) for img in images)
            return clean_text(ocr_text) or fast_text
        except Exception: return fast_text
    return ""

# ----- NEW: Secret Reset Route ----- #

@app.get("/api/v1/admin/reset")
async def manual_reset(key: str):
    """The Secret Reset API. Usage: /api/v1/admin/reset?key=ap_finance_2025"""
    # You can change this key to any secure string you prefer
    if key == "ap_finance_2025": 
        reset_system_data()
        return {"status": "success", "message": "Backend data wiped. Ready for fresh testing."}
    raise HTTPException(status_code=403, detail="Unauthorized: Invalid Secret Key")

# ----- Main Scan Routes (Keep existing) ----- #

@app.get("/api/v1/dashboard/stats")
def get_dashboard_stats():
    # Content remains identical to your original
    return {
        "summary": {
            "total_scanned": 14205,
            "fraud_detected": 45,
            "savings_in_crores": 1.2,
        },
        "weekly_activity": [
            {"day": "Mon", "uploads": 120, "fraud": 2},
            {"day": "Tue", "uploads": 150, "fraud": 5},
            {"day": "Wed", "uploads": 180, "fraud": 1},
            {"day": "Thu", "uploads": 90, "fraud": 0},
            {"day": "Fri", "uploads": 200, "fraud": 8},
            {"day": "Sat", "uploads": 50, "fraud": 0},
            {"day": "Sun", "uploads": 30, "fraud": 0},
        ],
        "recent_scans": [
            {"id": "1", "filename": "invoice_992.pdf", "status": "safe", "timestamp": "2 mins ago"},
            {"id": "2", "filename": "contract_v2.docx", "status": "warning", "timestamp": "5 mins ago"},
        ],
    }

@app.post("/api/v1/scan/upload")
async def upload_scan(file: UploadFile = File(...)):
    start_time = datetime.now()
    content = await file.read()
    filename = file.filename
    task_id = str(uuid.uuid4())

    text = extract_text_from_file(content, filename)
    entities = extract_advanced_entities(text)

    tables = []
    if filename.lower().endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                tables = [p.extract_table() for p in pdf.pages if p.extract_table()]
        except Exception: pass

    img_hash = get_image_phash(content)
    is_dup, dup_score = search_duplicate(text, img_hash)
    tamper_msg, tamper_conf = detect_tampering(content, filename)
    meta_issue, meta_conf = analyze_metadata(content, text)
    pii_found, pii_conf = detect_pii(text)

    fraud_score = 0
    anomalies = []

    if tamper_msg:
        fraud_score = max(fraud_score, 90)
        anomalies.append({"type": "Forensic Tampering", "description": tamper_msg, "confidence": tamper_conf})

    if meta_issue:
        fraud_score = max(fraud_score, 85)
        anomalies.append({"type": "Metadata Fraud", "description": meta_issue, "confidence": meta_conf})

    if is_dup:
        fraud_score = 100
        anomalies.append({"type": "Duplicate Discovery", "description": "Visual or text match found.", "confidence": dup_score})

    if pii_found:
        anomalies.append({"type": "PII Detected", "description": f"Contains: {pii_found}", "confidence": pii_conf})
        if fraud_score < 30: fraud_score += 20

    severity = "CRITICAL" if fraud_score >= 70 else "WARNING" if fraud_score >= 30 else "SAFE"
    
    if not is_dup:
        add_to_index(text, img_hash)

    overall_confidence = sum(a["confidence"] for a in anomalies) / len(anomalies) if anomalies else 0.0

    result = {
        "file_id": task_id,
        "filename": filename,
        "fraud_score": min(100, fraud_score),
        "severity": severity,
        "anomalies": anomalies,
        "text_content": text,
        "entities": entities,
        "extracted_tables": tables,
        "processing_time": int((datetime.now() - start_time).total_seconds() * 1000),
        "confidence": overall_confidence,
        "status": "completed"
    }

    db[task_id] = result
    return {"task_id": task_id, "message": "Unified fraud analysis concluded."}

@app.get("/api/v1/scan/result/{task_id}")
async def get_result(task_id: str):
    return db.get(task_id, {"error": "Verdict not found"})

@app.get("/health")
@app.get("/api/v1/health")
def health_check():
    return {"status": "healthy", "version": "1.0.0", "service": "AP FraudShield API"}

@app.post("/api/v1/admin/trigger-alert", response_model=AlertResponse)
def trigger_alert(payload: AlertRequest):
    return {"status": "sent"}